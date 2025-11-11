from __future__ import annotations

import json
import os
from dataclasses import dataclass
from typing import List, Optional, Dict, Any

from .audio_processor import ms_to_hhmmss, ms_to_srt_timestamp

try:
	from docx import Document  # type: ignore
	has_docx = True
except Exception:
	has_docx = False


@dataclass
class TranscriptSegment:
	start_ms: int
	end_ms: int
	text: str
	speaker: str = "Speaker 1"


def ensure_dir(path: str) -> None:
	if not os.path.exists(path):
		os.makedirs(path, exist_ok=True)


def clean_text(text: str) -> str:
	"""Clean text by removing non-printable and garbled characters."""
	# Remove zero-width spaces, non-breaking spaces, and other problematic characters
	text = text.replace('\u200b', '').replace('\u00a0', ' ').replace('\ufeff', '')
	# Remove any character that's not printable ASCII or common punctuation
	cleaned = ''.join(char for char in text if char.isprintable() or char in '\n\r\t')
	# Replace multiple spaces with single space
	cleaned = ' '.join(cleaned.split())
	cleaned = cleaned.strip()
	
	# Filter out segments that are only punctuation/question marks (Whisper artifacts)
	# Remove all punctuation and check if there's any actual text left
	text_only = ''.join(char for char in cleaned if char.isalnum() or char.isspace())
	if not text_only.strip():
		return ""  # No actual text content, return empty
	
	# Filter out segments that are mostly question marks (likely silence/noise)
	if cleaned.count('?') > len(cleaned) * 0.5:  # More than 50% question marks
		return ""
	
	return cleaned


def export_txt(segments: List[TranscriptSegment], out_dir: str, base_name: str) -> str:
	ensure_dir(out_dir)
	out_path = os.path.join(out_dir, f"{base_name}.txt")
	with open(out_path, "w", encoding="utf-8") as f:
		for seg in segments:
			start = ms_to_hhmmss(seg.start_ms)
			end = ms_to_hhmmss(seg.end_ms)
			cleaned_text = clean_text(seg.text)
			if cleaned_text:  # Only write if there's actual text content
				f.write(f"[{start} - {end}] {seg.speaker}: {cleaned_text}\n")
	return out_path


def export_txt_with_speakers(segments: List[TranscriptSegment], out_dir: str, base_name: str) -> str:
	"""Export TXT with AI-identified speaker names (creates _speakers.txt file)."""
	ensure_dir(out_dir)
	out_path = os.path.join(out_dir, f"{base_name}_speakers.txt")
	with open(out_path, "w", encoding="utf-8") as f:
		for seg in segments:
			start = ms_to_hhmmss(seg.start_ms)
			end = ms_to_hhmmss(seg.end_ms)
			cleaned_text = clean_text(seg.text)
			if cleaned_text:  # Only write if there's actual text content
				f.write(f"[{start} - {end}] {seg.speaker}: {cleaned_text}\n")
	return out_path


def export_json(segments: List[TranscriptSegment], out_dir: str, base_name: str, metadata: Optional[Dict[str, Any]] = None) -> str:
	ensure_dir(out_dir)
	out_path = os.path.join(out_dir, f"{base_name}.json")
	
	# Filter and clean segments
	cleaned_segments = []
	for s in segments:
		cleaned_text = clean_text(s.text)
		if cleaned_text:  # Only include segments with actual text
			cleaned_segments.append({
				"start_ms": s.start_ms,
				"end_ms": s.end_ms,
				"start": ms_to_hhmmss(s.start_ms),
				"end": ms_to_hhmmss(s.end_ms),
				"speaker": s.speaker,
				"text": cleaned_text,
			})
	
	payload = {
		"metadata": metadata or {},
		"segments": cleaned_segments,
	}
	
	# Ensure all metadata values are JSON serializable
	if metadata:
		serializable_metadata = {}
		for key, value in metadata.items():
			if callable(value):
				serializable_metadata[key] = str(value)
			else:
				serializable_metadata[key] = value
		payload["metadata"] = serializable_metadata
	
	with open(out_path, "w", encoding="utf-8") as f:
		json.dump(payload, f, ensure_ascii=False, indent=2)
	return out_path


def export_srt(segments: List[TranscriptSegment], out_dir: str, base_name: str) -> str:
	ensure_dir(out_dir)
	out_path = os.path.join(out_dir, f"{base_name}.srt")
	with open(out_path, "w", encoding="utf-8") as f:
		idx = 1
		for seg in segments:
			cleaned_text = clean_text(seg.text)
			if cleaned_text:  # Only include segments with actual text
				start = ms_to_srt_timestamp(seg.start_ms)
				end = ms_to_srt_timestamp(seg.end_ms)
				text = f"{seg.speaker}: {cleaned_text}".strip()
				f.write(f"{idx}\n{start} --> {end}\n{text}\n\n")
				idx += 1
	return out_path


def export_docx(segments: List[TranscriptSegment], out_dir: str, base_name: str) -> str:
	if not has_docx:
		raise RuntimeError("python-docx is not installed. Please install 'python-docx' to export DOCX.")
	ensure_dir(out_dir)
	out_path = os.path.join(out_dir, f"{base_name}.docx")
	doc = Document()
	doc.add_heading("Meeting Transcript", level=1)
	for seg in segments:
		cleaned_text = clean_text(seg.text)
		if cleaned_text:  # Only include segments with actual text
			start = ms_to_hhmmss(seg.start_ms)
			end = ms_to_hhmmss(seg.end_ms)
			doc.add_paragraph(f"[{start} - {end}] {seg.speaker}: {cleaned_text}")
	doc.save(out_path)
	return out_path


